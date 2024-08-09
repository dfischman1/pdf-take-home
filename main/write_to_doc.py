from docx import Document



def write_out_docx(q_and_a: dict):
    document = Document()
    document.add_heading('Questions and Answers for Medical Record', 0)
    for question, answer in q_and_a.items():
        document.add_paragraph(question)
        document.add_paragraph(answer)
        document.add_page_break()
    document.save("questions_and_answers_about_medical_record.docx")


def write_out_txt(q_and_a: dict, file_name: str):
    with open(f"{file_name}.txt", "w") as file:
        for question, answer in q_and_a.items():
            file.write(question + "\n")
            file.write(answer)
            file.write("\n\n")